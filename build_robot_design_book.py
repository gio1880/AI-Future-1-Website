from pathlib import Path
import math
import re
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Robot Design Technical Book"
ASSETS = OUT / "generated_assets"
DOCX = OUT / "AI Future - Robot Design Technical Book.docx"
PDF = OUT / "AI Future - Robot Design Technical Book.pdf"
PROOF = OUT / "pdf_page_proofs"

NAVY = "#0B2545"
INK = "#111827"
MUTED = "#64748B"
LIGHT = "#F4F7FA"
BORDER = "#D9E2EC"
GREEN = "#2F9E6D"
GOLD = "#C28A13"
RED = "#C44545"

RUNS = [
    ("Run A", "M5 M6 M7 M8", "#E84A4A", "right launch", "Push/lift chain with fast return"),
    ("Run B", "M9 M10", "#2F80ED", "right launch", "Tip and pan control"),
    ("Run C", "M1 M2", "#D89C16", "left-side route", "Pickup geometry"),
    ("Run D", "M3 M4", "#8E44AD", "left launch", "Precision lift and drop"),
    ("Run E", "M11 M13", "#00A6A6", "right launch", "Long blue-side transition"),
    ("Run G", "M15", "#6C63FF", "left launch", "Slow controlled final drop"),
]

MISSION_POINTS = {
    "Run A": [(724, 410), (724, 300), (674, 185), (650, 91), (603, 67), (692, 69), (730, 128), (714, 219), (724, 410)],
    "Run B": [(724, 410), (562, 212), (646, 309), (724, 410)],
    "Run C": [(61, 160), (88, 46), (61, 160)],
    "Run D": [(56, 389), (89, 358), (106, 313), (105, 254), (108, 43), (223, 29), (154, 105), (109, 235), (56, 389)],
    "Run E": [(724, 410), (656, 408), (589, 397), (512, 380), (446, 343), (393, 276), (346, 207), (316, 178), (295, 198)],
    "Run G": [(56, 389), (94, 260), (56, 389)],
}

MISSION_LABELS = {
    "Run A": [("5", 603, 67), ("6", 650, 91), ("7", 692, 69), ("8", 714, 219)],
    "Run B": [("9", 562, 212), ("10", 646, 309)],
    "Run C": [("1", 61, 160), ("2", 88, 46)],
    "Run D": [("3", 223, 29), ("4", 154, 105)],
    "Run E": [("11", 512, 380), ("13", 316, 178)],
    "Run G": [("15", 94, 260)],
}

ATTACHMENTS = {
    "Run A": [
        "Attachments/Run A/run a v1.jpg",
        "Attachments/Run A/run a v2.jpg",
        "Attachments/Run A/Run A attachment V2.jpg",
        "Attachments/Run A/run a iteration 4.jpg",
    ],
    "Run B": [
        "Attachments/Run B/run b v1.jpg",
        "Attachments/Run B/run b v2.jpg",
        "Attachments/Run B/run b attachment v2.jpg",
        "Attachments/Run B/run b iteration 4.jpg",
    ],
    "Run C": [
        "Attachments/Run C/run c attachment v1.jpg",
        "Attachments/Run C/run c attachment v2.jpg",
        "Attachments/Run C/run c v3.1.jpg",
        "Attachments/Run C/run c iteration 4.jpg",
    ],
    "Run D": [
        "Attachments/Run D/run d attachment v1.jpg",
        "Attachments/Run D/run d attachment v1.jpg",
        "Attachments/Run D/run d iteration 4.JPG",
        "Attachments/Run D/run d iteration 4.JPG",
    ],
    "Run E": [
        "Attachments/Run E/run e v2.jpg",
        "Attachments/Run E/run e v2.1.jpg",
        "Attachments/Run E/run 3 attachment V1.jpg",
        "Attachments/Run E/run e v3.jpg",
    ],
}


def hex_to_rgb(hex_color):
    h = hex_color.lstrip("#")
    return tuple(int(h[i:i + 2], 16) for i in (0, 2, 4))


def font(size=28, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def cover_fit(path, size):
    img = Image.open(path).convert("RGB")
    return ImageOps.fit(img, size, method=Image.Resampling.LANCZOS, centering=(0.5, 0.5))


def contain_fit(path, size, fill=(255, 255, 255)):
    try:
        img = Image.open(path).convert("RGB")
    except Exception:
        img = Image.new("RGB", size, fill)
        d = ImageDraw.Draw(img)
        d.text((20, 20), "Image unavailable", fill=hex_to_rgb(MUTED), font=font(22, True))
        return img
    img.thumbnail(size, Image.Resampling.LANCZOS)
    canvas = Image.new("RGB", size, fill)
    canvas.paste(img, ((size[0] - img.width) // 2, (size[1] - img.height) // 2))
    return canvas


def draw_wrapped(draw, xy, text, fnt, fill, width, line_gap=6):
    x, y = xy
    avg = max(1, int(width / (fnt.size * 0.55)))
    lines = []
    for para in str(text).splitlines():
        lines.extend(textwrap.wrap(para, avg) or [""])
    for line in lines:
        draw.text((x, y), line, fill=fill, font=fnt)
        y += fnt.size + line_gap
    return y


def round_rect(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def save_svg_placeholder(name, title, body, color=NAVY):
    svg = f'''<svg xmlns="http://www.w3.org/2000/svg" width="1200" height="720" viewBox="0 0 1200 720">
  <rect width="1200" height="720" fill="#F7FAFC"/>
  <rect x="48" y="48" width="1104" height="624" rx="28" fill="#FFFFFF" stroke="#D9E2EC" stroke-width="3"/>
  <text x="80" y="120" font-family="Arial" font-size="52" font-weight="700" fill="{color}">{title}</text>
  <text x="80" y="180" font-family="Arial" font-size="28" fill="#64748B">{body}</text>
</svg>'''
    (ASSETS / f"{name}.svg").write_text(svg, encoding="utf-8")


def create_cover():
    img_path = ROOT / "Robot Iterations/base robot iteration 6.2.jpg"
    bg = cover_fit(img_path, (1600, 2100))
    overlay = Image.new("RGBA", bg.size, (255, 255, 255, 0))
    d = ImageDraw.Draw(overlay)
    d.rectangle((0, 0, 1600, 2100), fill=(255, 255, 255, 72))
    d.rectangle((0, 1320, 1600, 2100), fill=(255, 255, 255, 232))
    d.rectangle((0, 0, 1600, 120), fill=hex_to_rgb(NAVY) + (235,))
    d.text((90, 38), "AI FUTURE  |  UNEARTHED SEASON", fill="white", font=font(38, True))
    d.text((90, 1410), "Robot Design", fill=hex_to_rgb(NAVY), font=font(105, True))
    d.text((90, 1532), "& Mission Strategy", fill=hex_to_rgb(NAVY), font=font(88, True))
    d.text((94, 1655), "Technical book for judges, table display, and team presentation", fill=hex_to_rgb(MUTED), font=font(36))
    for i, (run, missions, color, _, _) in enumerate(RUNS):
        x = 94 + i * 235
        round_rect(d, (x, 1785, x + 200, 1865), 18, hex_to_rgb(color) + (235,))
        d.text((x + 18, 1808), run, fill="white", font=font(27, True))
    out = ASSETS / "cover.png"
    Image.alpha_composite(bg.convert("RGBA"), overlay).convert("RGB").save(out, quality=94)
    return out


def create_robot_anatomy():
    base = contain_fit(ROOT / "Robot Iterations/base robot iteration 6.2.jpg", (1400, 900), fill=(248, 250, 252))
    d = ImageDraw.Draw(base, "RGBA")
    d.rectangle((0, 0, 1400, 900), outline=hex_to_rgb(BORDER), width=4)
    labels = [
        ((80, 90), (520, 250), "Compact base", "Lower, wider base reduces wobble."),
        ((930, 90), (810, 270), "SPIKE Prime hub", "One shared hub initialized through robot_base.py."),
        ((80, 665), (540, 580), "Drive motors", "Ports A/E drive the calibrated DriveBase."),
        ((920, 665), (760, 590), "Attachment motors", "LAM/RAM power active mission mechanisms."),
    ]
    for box_xy, anchor, title, body in labels:
        x, y = box_xy
        round_rect(d, (x, y, x + 380, y + 122), 18, (255, 255, 255, 235), hex_to_rgb(BORDER), 2)
        d.text((x + 22, y + 18), title, fill=hex_to_rgb(NAVY), font=font(30, True))
        draw_wrapped(d, (x + 22, y + 58), body, font(22), hex_to_rgb(INK), 328)
        d.line((x + 190, y + 122, anchor[0], anchor[1]), fill=hex_to_rgb(NAVY) + (210,), width=5)
        d.ellipse((anchor[0] - 9, anchor[1] - 9, anchor[0] + 9, anchor[1] + 9), fill=hex_to_rgb(NAVY) + (240,))
    out = ASSETS / "robot_anatomy_overlay.png"
    base.save(out, quality=94)
    save_svg_placeholder("robot_anatomy_overlay", "Robot Anatomy Overlay", "Label source for generated PNG overlay.")
    return out


def create_mission_strategy():
    canvas_w, canvas_h = 1450, 1800
    map_x, map_y, map_w, map_h = 75, 210, 1300, 760
    base = Image.new("RGB", (canvas_w, canvas_h), "white")
    d = ImageDraw.Draw(base, "RGBA")
    d.rectangle((0, 0, canvas_w, canvas_h), fill=(248, 250, 252, 255))
    d.text((70, 55), "Mission Run Map", fill=hex_to_rgb(NAVY), font=font(58, True))
    d.text((70, 122), "Each colored path shows one planned run group on the Unearthed field.", fill=hex_to_rgb(MUTED), font=font(30))
    mat = Image.open(ROOT / "Mission Pictures/fll unearthed map.png").convert("RGB").resize((map_w, map_h), Image.Resampling.LANCZOS)
    base.paste(mat, (map_x, map_y))
    d.rectangle((map_x, map_y, map_x + map_w, map_y + map_h), outline=hex_to_rgb(NAVY), width=6)
    scale_x = map_w / 798
    scale_y = map_h / 466
    for run, missions, color, _, _ in RUNS:
        pts = [(map_x + int(x * scale_x), map_y + int(y * scale_y)) for x, y in MISSION_POINTS[run]]
        for offset in [14, 8]:
            d.line(pts, fill=hex_to_rgb(color) + (55,), width=offset)
        d.line(pts, fill=hex_to_rgb(color) + (235,), width=7)
        for label, x, y in MISSION_LABELS[run]:
            cx, cy = map_x + int(x * scale_x), map_y + int(y * scale_y)
            d.ellipse((cx - 31, cy - 31, cx + 31, cy + 31), fill=hex_to_rgb(color) + (245,), outline=(255, 255, 255, 255), width=5)
            tw = d.textlength(label, font=font(28, True))
            d.text((cx - tw / 2, cy - 18), label, fill="white", font=font(28, True))
    y = 1030
    for idx, (run, missions, color, launch, note) in enumerate(RUNS):
        col = idx % 2
        row = idx // 2
        x = 75 + col * 665
        by = y + row * 220
        round_rect(d, (x, by, x + 620, by + 175), 24, (255, 255, 255, 255), hex_to_rgb(color), 4)
        d.rectangle((x, by, x + 34, by + 175), fill=hex_to_rgb(color))
        d.text((x + 58, by + 28), f"{run}: {missions}", fill=hex_to_rgb(NAVY), font=font(31, True))
        d.text((x + 58, by + 75), f"Start: {launch}", fill=hex_to_rgb(MUTED), font=font(23, True))
        draw_wrapped(d, (x + 58, by + 108), note, font(23), hex_to_rgb(INK), 520, line_gap=3)
    out = ASSETS / "mission_strategy_overlay.png"
    base.save(out, quality=94)
    svg_lines = []
    for run, missions, color, _, _ in RUNS:
        pts = " ".join(f"{map_x + x * scale_x:.1f},{map_y + y * scale_y:.1f}" for x, y in MISSION_POINTS[run])
        svg_lines.append(f'<polyline points="{pts}" fill="none" stroke="{color}" stroke-width="7" stroke-linecap="round" stroke-linejoin="round" opacity="0.92"/>')
        for label, x, y in MISSION_LABELS[run]:
            cx, cy = map_x + x * scale_x, map_y + y * scale_y
            svg_lines.append(f'<circle cx="{cx:.1f}" cy="{cy:.1f}" r="31" fill="{color}" stroke="#FFFFFF" stroke-width="5"/>')
            svg_lines.append(f'<text x="{cx:.1f}" y="{cy + 10:.1f}" text-anchor="middle" font-family="Arial" font-size="28" font-weight="700" fill="#FFFFFF">{label}</text>')
    legend = []
    for idx, (run, missions, color, launch, note) in enumerate(RUNS):
        col = idx % 2
        row = idx // 2
        x = 75 + col * 665
        by = 1030 + row * 220
        legend.append(f'<rect x="{x}" y="{by}" width="620" height="175" rx="24" fill="#FFFFFF" stroke="{color}" stroke-width="4"/>')
        legend.append(f'<rect x="{x}" y="{by}" width="34" height="175" fill="{color}"/>')
        legend.append(f'<text x="{x + 58}" y="{by + 62}" font-family="Arial" font-size="31" font-weight="700" fill="{NAVY}">{run}: {missions}</text>')
        legend.append(f'<text x="{x + 58}" y="{by + 105}" font-family="Arial" font-size="23" font-weight="700" fill="{MUTED}">Start: {launch}</text>')
        legend.append(f'<text x="{x + 58}" y="{by + 140}" font-family="Arial" font-size="23" fill="{INK}">{note}</text>')
    svg = f'''<svg xmlns="http://www.w3.org/2000/svg" width="{canvas_w}" height="{canvas_h}" viewBox="0 0 {canvas_w} {canvas_h}">
  <rect width="{canvas_w}" height="{canvas_h}" fill="#F8FAFC"/>
  <text x="70" y="100" font-family="Arial" font-size="58" font-weight="700" fill="{NAVY}">Mission Run Map</text>
  <text x="70" y="150" font-family="Arial" font-size="30" fill="{MUTED}">Each colored path shows one planned run group on the Unearthed field.</text>
  <image x="{map_x}" y="{map_y}" width="{map_w}" height="{map_h}" href="../../Mission%20Pictures/fll%20unearthed%20map.png" preserveAspectRatio="none"/>
  <rect x="{map_x}" y="{map_y}" width="{map_w}" height="{map_h}" fill="none" stroke="{NAVY}" stroke-width="6"/>
  {''.join(svg_lines)}
  {''.join(legend)}
</svg>'''
    (ASSETS / "mission_strategy_overlay.svg").write_text(svg, encoding="utf-8")
    return out


def create_attachment_timeline(run, color):
    paths = ATTACHMENTS[run]
    labels = ["Prototype", "Improve", "Refine", "Final"]
    outputs = []
    for page_idx, start in enumerate((0, 2), start=1):
        canvas = Image.new("RGB", (1450, 1800), "white")
        d = ImageDraw.Draw(canvas, "RGBA")
        d.rectangle((0, 0, 1450, 1800), fill=(248, 250, 252, 255))
        d.text((48, 38), f"{run} Attachment Iteration", fill=hex_to_rgb(NAVY), font=font(48, True))
        d.text((48, 100), f"Vertical evidence page {page_idx}: versions {start + 1}-{start + 2}", fill=hex_to_rgb(MUTED), font=font(28))
        y = 175
        for i in range(start, start + 2):
            rel = paths[i]
            round_rect(d, (48, y, 1402, y + 720), 28, (255, 255, 255, 255), hex_to_rgb(BORDER), 3)
            d.rectangle((48, y, 1402, y + 78), fill=hex_to_rgb(color) + (235,))
            d.text((78, y + 22), f"{i + 1}. {labels[i]}", fill="white", font=font(33, True))
            photo = contain_fit(ROOT / rel, (1294, 545), fill=(255, 255, 255))
            canvas.paste(photo, (78, y + 108))
            d.text((78, y + 668), f"Evidence photo v{i + 1}", fill=hex_to_rgb(NAVY), font=font(27, True))
            y += 765
        d.text((48, 1730), "Design rule: change one main variable at a time so the team can tell what improved reliability.", fill=hex_to_rgb(MUTED), font=font(25, True))
        out = ASSETS / f"{run.lower().replace(' ', '_')}_attachment_timeline_{page_idx}.png"
        canvas.save(out, quality=94)
        outputs.append(out)
    save_svg_placeholder(f"{run.lower().replace(' ', '_')}_attachment_timeline", f"{run} Vertical Attachment Timeline", "Generated from local attachment photos.", color)
    return outputs


def create_code_architecture():
    img = Image.new("RGB", (1450, 820), "white")
    d = ImageDraw.Draw(img, "RGBA")
    d.rectangle((0, 0, 1450, 820), fill=(248, 250, 252, 255))
    d.text((50, 42), "Programming architecture", fill=hex_to_rgb(NAVY), font=font(46, True))
    d.text((50, 100), "Menu selection imports one run file, while robot_base.py keeps hardware setup consistent.", fill=hex_to_rgb(MUTED), font=font(25))
    boxes = [
        ((70, 190, 390, 315), "Menu.py", "Hub menu: 1-6 or D"),
        ((535, 190, 915, 315), "Run files", "Run_A, Run_B, Run_C, Run_D, Run_E, Run_G"),
        ((1055, 190, 1375, 315), "diagnostics.py", "Health checks before runs"),
        ((535, 455, 915, 610), "robot_base.py", "PrimeHub, motors, DriveBase, gyro reset"),
        ((70, 650, 390, 760), "Hardware", "A/E drive motors, B/C attachment motors"),
        ((1055, 650, 1375, 760), "Control", "Gyro heading + encoder distance"),
    ]
    for box, title, body in boxes:
        fill = (255, 255, 255, 255)
        outline = hex_to_rgb(NAVY) if title in ("Menu.py", "robot_base.py") else hex_to_rgb(BORDER)
        round_rect(d, box, 20, fill, outline, 3)
        d.text((box[0] + 24, box[1] + 22), title, fill=hex_to_rgb(NAVY), font=font(30, True))
        draw_wrapped(d, (box[0] + 24, box[1] + 64), body, font(22), hex_to_rgb(INK), box[2] - box[0] - 48)
    for line in [
        (390, 252, 535, 252), (915, 252, 1055, 252), (725, 315, 725, 455),
        (535, 532, 390, 700), (915, 532, 1055, 700),
    ]:
        d.line(line, fill=hex_to_rgb(GREEN) + (230,), width=6)
    out = ASSETS / "code_architecture_overlay.png"
    img.save(out, quality=94)
    save_svg_placeholder("code_architecture_overlay", "Code Architecture", "Menu, run files, shared base, diagnostics.")
    return out


def create_diagnostics_dashboard():
    img = Image.new("RGB", (1450, 760), "white")
    d = ImageDraw.Draw(img, "RGBA")
    d.rectangle((0, 0, 1450, 760), fill=(248, 250, 252, 255))
    d.text((50, 42), "Competition readiness dashboard", fill=hex_to_rgb(NAVY), font=font(46, True))
    checks = [
        ("Gyro Drift", "Max z-axis drift < 0.5 deg/s and heading stays near zero.", GREEN),
        ("Motor Sync", "Compare left/right encoder angles after a 2-second run.", GREEN),
        ("Turn Accuracy", "Four 90-degree turns should finish near 360 degrees.", GREEN),
        ("Battery Voltage", "Minimum 7400 mV to avoid weak, inconsistent motion.", GOLD),
        ("Status Lights", "Yellow running, green pass, red fail for quick table checks.", NAVY),
    ]
    x, y = 55, 140
    for i, (title, body, color) in enumerate(checks):
        col = i % 2
        row = i // 2
        bx = x + col * 675
        by = y + row * 175
        round_rect(d, (bx, by, bx + 620, by + 135), 20, (255, 255, 255, 255), hex_to_rgb(BORDER), 2)
        d.ellipse((bx + 26, by + 34, bx + 76, by + 84), fill=hex_to_rgb(color) + (235,))
        d.text((bx + 104, by + 28), title, fill=hex_to_rgb(NAVY), font=font(29, True))
        draw_wrapped(d, (bx + 104, by + 68), body, font(21), hex_to_rgb(INK), 460)
    round_rect(d, (55, 642, 1395, 724), 18, (255, 255, 255, 255), hex_to_rgb(GREEN), 3)
    draw_wrapped(
        d,
        (82, 657),
        "Why this matters: diagnostics separate code tuning from hardware issues like loose cables, weak motors, low battery, or gyro drift.",
        font(22, True),
        hex_to_rgb(INK),
        1260,
        line_gap=4,
    )
    out = ASSETS / "diagnostics_dashboard_overlay.png"
    img.save(out, quality=94)
    save_svg_placeholder("diagnostics_dashboard_overlay", "Diagnostics Dashboard", "Gyro drift, motor sync, turns, battery, lights.")
    return out


def create_mission_grid():
    img = Image.new("RGB", (1450, 1000), "white")
    d = ImageDraw.Draw(img, "RGBA")
    d.rectangle((0, 0, 1450, 1000), fill=(248, 250, 252, 255))
    d.text((42, 30), "Mission model grouping", fill=hex_to_rgb(NAVY), font=font(44, True))
    y = 105
    for run, missions, color, launch, note in RUNS:
        round_rect(d, (38, y, 1410, y + 132), 18, (255, 255, 255, 255), hex_to_rgb(BORDER), 2)
        d.rectangle((38, y, 168, y + 132), fill=hex_to_rgb(color) + (235,))
        d.text((56, y + 35), run, fill="white", font=font(26, True))
        x = 190
        for mission in missions.split():
            number = re.sub(r"\D", "", mission)
            photo = contain_fit(ROOT / f"Mission Pictures/mission {number}.png", (92, 92), fill=(255, 255, 255))
            img.paste(photo, (x, y + 20))
            d.text((x + 32, y + 105), mission, fill=hex_to_rgb(NAVY), font=font(16, True))
            x += 112
        d.text((780, y + 28), f"Start: {launch}", fill=hex_to_rgb(NAVY), font=font(24, True))
        d.text((780, y + 66), note, fill=hex_to_rgb(INK), font=font(22))
        y += 145
    out = ASSETS / "mission_model_grid.png"
    img.save(out, quality=94)
    return out


def create_assets():
    OUT.mkdir(exist_ok=True)
    ASSETS.mkdir(exist_ok=True)
    made = {
        "cover": create_cover(),
        "robot": create_robot_anatomy(),
        "mission": create_mission_strategy(),
        "code": create_code_architecture(),
        "diagnostics": create_diagnostics_dashboard(),
        "mission_grid": create_mission_grid(),
    }
    for run, _, color, _, _ in RUNS:
        if run in ATTACHMENTS:
            made[f"{run}_timeline"] = create_attachment_timeline(run, color)
    return made


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.lstrip("#"))
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color.lstrip("#"))


def add_run(run, text, size=11, bold=False, color=INK):
    r = text.add_run() if False else None


def style_run(run, size=11, bold=False, color=INK, font_name="Aptos"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(*hex_to_rgb(color))


def add_p(doc, text="", style=None, size=10.5, color=INK, bold=False, align=None, after=4):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        style_run(r, size=size, bold=bold, color=color)
    return p


def add_h(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    style_run(r, size=18 if level == 1 else 13, bold=True, color=NAVY if level == 1 else "#1F4D78")
    return p


def add_label(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    style_run(r, size=8.5, bold=True, color=color)
    return p


def add_picture(doc, path, width=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(width))
    return p


def add_caption(doc, text):
    return add_p(doc, text, size=8.2, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)


def add_cell_picture(cell, path, width):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(path), width=Inches(width))


def add_cell_text(cell, text, size=9.2, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=color)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        r = p.add_run(item)
        style_run(r, size=10.1, color=INK)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        style_run(r, size=8.8, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            style_run(r, size=8.7, color=INK)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def set_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("AI Future - Robot Design Technical Book")
    style_run(r, size=8.2, color=MUTED)


def add_page_title(doc, tag, title, summary=None):
    add_label(doc, tag)
    add_h(doc, title, 1)
    if summary:
        add_p(doc, summary, size=10.4, color=MUTED, after=5)


def read_code(path):
    text = (ROOT / path).read_text(encoding="utf-8", errors="replace")
    text = text.replace("â€”", "-").replace("→", "->")
    return text


def code_excerpt(text, contains, before=1, after=8):
    lines = text.splitlines()
    idx = next((i for i, line in enumerate(lines) if contains in line), 0)
    start = max(0, idx - before)
    end = min(len(lines), idx + after)
    return "\n".join(lines[start:end])


def add_code_box(doc, title, code, note):
    add_label(doc, title, "#2563EB")
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.45)
    set_cell_shading(cell, "F8FAFC")
    set_cell_border(cell, "CBD5E1")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(code)
    style_run(r, size=7.8, color="#0F172A", font_name="Consolas")
    add_p(doc, note, size=8.5, color=MUTED, after=4)


def add_editable_mission_strategy(doc):
    add_p(doc, "This DOCX section uses editable Word elements: the run cards, labels, and strategy notes can be changed directly in Word.", size=9.2, color=MUTED)
    add_picture(doc, ROOT / "Mission Pictures/fll unearthed map.png", width=7.05)
    add_caption(doc, "Mission mat image with editable run table below. Use the generated SVG/PDF proof for the path overlay version.")
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["Run", "Missions", "Launch", "Design focus", "Mission models"]
    widths = [0.65, 1.05, 1.0, 2.1, 2.15]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(cell)
        add_cell_text(cell, header, size=8.3, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for run, missions, color, launch, note in RUNS:
        row = table.add_row().cells
        for i, w in enumerate(widths):
            row[i].width = Inches(w)
            set_cell_border(row[i])
            row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(row[0], color)
        add_cell_text(row[0], run, size=8.8, color="#FFFFFF", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(row[1], missions, size=8.8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(row[2], launch, size=8.3, color=MUTED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(row[3], note, size=8.4)
        row[4].text = ""
        p = row[4].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for mission in missions.split():
            number = re.sub(r"\D", "", mission)
            pic = ROOT / f"Mission Pictures/mission {number}.png"
            if pic.exists():
                p.add_run().add_picture(str(pic), width=Inches(0.34))
                p.add_run(" ")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_editable_attachment_pages(doc, run, color, takeaway):
    labels = ["Prototype", "Improve", "Refine", "Final"]
    paths = ATTACHMENTS[run]
    for page_idx, start in enumerate((0, 2), start=1):
        page_note = "versions 1-2" if page_idx == 1 else "versions 3-4"
        add_page_title(doc, "Attachment iteration", f"{run}: {page_note}", takeaway if page_idx == 1 else "Continuation of the vertical evidence timeline.")
        add_p(doc, "Editable Word layout: each label, caption, and note is real text; each photo is inserted separately so it can be resized or replaced.", size=8.8, color=MUTED)
        for i in range(start, start + 2):
            table = doc.add_table(rows=2, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            title_cell = table.cell(0, 0)
            photo_cell = table.cell(1, 0)
            title_cell.width = Inches(7.05)
            photo_cell.width = Inches(7.05)
            set_cell_shading(title_cell, color)
            set_cell_border(title_cell, color)
            set_cell_border(photo_cell, "CBD5E1")
            add_cell_text(title_cell, f"{i + 1}. {labels[i]}", size=11, color="#FFFFFF", bold=True)
            photo_cell.text = ""
            p = photo_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(ROOT / paths[i]), width=Inches(6.55))
            p2 = photo_cell.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            r = p2.add_run(f"Evidence photo v{i + 1}")
            style_run(r, size=9.5, bold=True, color=NAVY)
            doc.add_paragraph().paragraph_format.space_after = Pt(5)
        add_p(doc, "Design rule: change one main variable at a time so the team can tell what improved reliability.", size=9.5, color=MUTED, bold=True)
        doc.add_page_break()


def add_editable_code_architecture(doc):
    add_p(doc, "Editable Word architecture: change any file name, note, or arrow label in the table below.", size=8.8, color=MUTED)
    add_table(doc, ["Layer", "Editable element", "What it does"], [
        ("Selection", "Menu.py", "Hub menu chooses run 1-6 or D for diagnostics."),
        ("Run logic", "Run_A, Run_B, Run_C, Run_D, Run_E, Run_G", "Each selected file contains the movement and attachment timing for one mission group."),
        ("Shared base", "robot_base.py", "Initializes PrimeHub, drive motors, attachment motors, DriveBase, gyro use, and safe reset."),
        ("Hardware", "Ports A/E, B/C", "A/E are drive motors; B/C are LAM/RAM attachment motors."),
        ("Control", "Gyro + encoders + PID", "Gyro handles heading, encoders measure distance, PID corrects drift."),
        ("Health checks", "diagnostics.py", "Runs gyro drift, motor sync, turn accuracy, and battery tests before competition runs."),
    ], widths=[1.05, 2.05, 3.25])


def add_editable_diagnostics(doc):
    add_p(doc, "Editable Word dashboard: each check is a table row you can change for judging notes.", size=8.8, color=MUTED)
    add_table(doc, ["Check", "Pass target", "Why judges should care"], [
        ("Gyro Drift", "Max z-axis drift < 0.5 deg/s; heading stays near zero.", "Confirms the robot is not turning in software while sitting still."),
        ("Motor Sync", "Left/right encoder difference <= 15 degrees.", "Catches weak motors, loose cables, worn gearing, or bad ports."),
        ("Turn Accuracy", "Four 90-degree turns finish near 360 degrees.", "Confirms axle track and gyro calibration before match runs."),
        ("Battery Voltage", "At least 7400 mV.", "Low battery changes torque and makes runs inconsistent."),
        ("Status Lights", "Yellow running, green pass, red fail.", "Fast visual feedback at the table before a match."),
    ], widths=[1.25, 2.55, 2.55])


def build_docx(assets):
    doc = Document()
    set_page(doc)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Normal", "List Bullet"]:
        r_pr = styles[style_name]._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), "Aptos")

    # Cover
    add_picture(doc, assets["cover"], width=7.15)
    add_p(doc, "Built from the team's website, robot code, mission model photos, and attachment iterations.", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # Page 1
    add_page_title(doc, "Judge quick summary", "What we engineered this season", "Our design goal was not only to score points, but to make a robot that the whole team could align, repair, explain, and trust under competition pressure.")
    add_bullets(doc, [
        "Stable base: lower, cleaner, and easier to inspect than earlier concepts.",
        "Mission strategy: grouped nearby missions into short runs to reduce risk.",
        "Attachment system: active + active + passive mechanisms for strong multi-mission tools.",
        "Programming: Pybricks with gyro-enabled DriveBase, encoder distance, and tuned PID.",
        "Competition readiness: a diagnostics menu checks gyro drift, motor sync, turn accuracy, and battery voltage.",
    ])
    add_table(doc, ["Engineering area", "Evidence judges can inspect"], [
        ("Base robot", "Current build photo, robot anatomy overlay, wheel/axle constants in code"),
        ("Strategy", "Run grouping board, mission model photos, colored field paths"),
        ("Attachments", "Prototype-to-final photo timelines for Runs A-E"),
        ("Code", "Menu selector, shared robot_base.py, per-run motion files"),
        ("Testing", "Diagnostics dashboard and consistency tuning narrative"),
    ], widths=[1.6, 4.75])
    doc.add_page_break()

    # Page 2
    add_page_title(doc, "Robot evolution", "From concept to a reliable competition base")
    add_p(doc, "The robot started as a design challenge: build something strong enough to carry attachments but simple enough for fast match setup. We moved toward a compact base, cleaner motor layout, and attachment access that teammates could understand quickly.", size=10.4)
    add_picture(doc, ROOT / "initial concept this year.png", width=3.25)
    add_caption(doc, "Early concept sketch from this season.")
    doc.add_page_break()

    # Page 3
    add_page_title(doc, "Current robot", "Anatomy of the final base")
    add_picture(doc, assets["robot"], width=7.15)
    add_caption(doc, "Overlay labels connect the physical robot to the code architecture and attachment system.")
    add_bullets(doc, [
        "Drive motors are initialized once in the shared base file.",
        "Attachment motors are separated into left and right mechanisms.",
        "The base is designed around repeatable launch alignment and fast swaps.",
    ])
    doc.add_page_break()

    # Page 4
    add_page_title(doc, "Mission strategy", "Grouped runs reduce risk")
    add_editable_mission_strategy(doc)
    doc.add_page_break()

    # Page 5
    add_page_title(doc, "Mission models", "What each run is built to solve")
    add_picture(doc, assets["mission_grid"], width=7.15)
    add_caption(doc, "Mission model photos are grouped by run so judges can connect strategy to the table.")
    doc.add_page_break()

    # Page 6
    add_page_title(doc, "Run plan", "Current competition map")
    add_table(doc, ["Run", "Missions", "Launch", "Design focus"], [(r, m, launch, note) for r, m, _, launch, note in RUNS], widths=[0.75, 1.2, 1.45, 3.0])
    add_p(doc, "The plan favors shorter trips, repeatable alignment, and attachments that solve multiple models before returning home.", size=10.4)
    add_bullets(doc, [
        "Run A handles the largest red-side group, so it received custom axle-track tuning and heading PID.",
        "Run D uses slower precision settings because lifting and dropping require accuracy more than speed.",
        "Run E crosses farther across the table, so it uses measured turns and print statements to check heading during tuning.",
    ])
    doc.add_page_break()

    # Pages 7-11
    takeaways = {
        "Run A": "Passive mechanism progression for cleaner trigger flow.",
        "Run B": "Tip control evolution for repeatable pan contact.",
        "Run C": "Pickup geometry tuned for more reliable retrieval.",
        "Run D": "Lift geometry evolved toward stable drop positioning.",
        "Run E": "Underside approach refined for faster transitions.",
    }
    for run in ATTACHMENTS:
        color = next(c for r, _, c, _, _ in RUNS if r == run)
        add_editable_attachment_pages(doc, run, color, takeaways[run])

    # Page 12
    add_page_title(doc, "Attachment system", "Two active actions plus one passive action")
    add_p(doc, "Our strongest attachments combined powered control with simple passive motion. The active mechanisms gave us precision, while passive geometry made setup and release faster.", size=10.4)
    add_table(doc, ["Design rule", "Why it helped"], [
        ("Active attachment", "Precise movement for lifts, drops, and controlled pushes."),
        ("Active attachment", "Second powered action lets one run solve more than one model."),
        ("Passive attachment", "No motor timing needed, so the action is faster and simpler."),
        ("Color coding", "Teammates can grab the right run tool instantly."),
    ], widths=[1.6, 4.75])
    add_bullets(doc, [
        "Run A: all black pieces.",
        "Run B: blue, grey, and yellow pieces.",
        "Run C: black and yellow pieces.",
        "Run D: all black pieces.",
        "Run E: all blue pieces.",
    ])
    doc.add_page_break()

    # Page 13
    add_page_title(doc, "Programming", "Pybricks gives us direct control")
    add_editable_code_architecture(doc)
    robot_base = read_code("Team Robot Code/robot_base.py")
    add_code_box(doc, "robot_base.py excerpt", code_excerpt(robot_base, "def make_robot", 0, 10), "The shared factory enables gyro, resets heading safely, and returns a DriveBase ready for each run.")
    doc.add_page_break()

    # Page 14
    add_page_title(doc, "Code decisions", "Real run files, real calibration choices")
    run_a = read_code("Team Robot Code/Run_A_updated.py")
    run_d = read_code("Team Robot Code/Run_D.py")
    add_code_box(doc, "Run_A_updated.py", code_excerpt(run_a, "robot.heading_control.pid", 3, 10), "Run A uses axle_track=101.6 and a tuned heading PID because it solves a large mission chain.")
    add_code_box(doc, "Run_D.py", code_excerpt(run_d, "robot.settings", 1, 8), "Run D uses axle_track=152 and slower settings for precision lift/drop work.")
    doc.add_page_break()

    # Page 15
    add_page_title(doc, "Menu and match flow", "One simple selector for competition")
    menu = read_code("Team Robot Code/Menu.py")
    add_code_box(doc, "Menu.py", code_excerpt(menu, "selected = hub_menu", 0, 18), "The center button becomes the stop button, the hub light gives status, and D opens diagnostics.")
    add_bullets(doc, [
        "The menu prevents hunting through files at the table.",
        "Each selected number imports exactly one run file.",
        "Diagnostics are available from the same interface as the mission runs.",
    ])
    doc.add_page_break()

    # Page 16
    add_page_title(doc, "Sensors and PID", "Gyro heading plus encoder distance")
    add_p(doc, "The robot uses the gyro for heading and the motor encoders for distance. Pybricks DriveBase lets us tune straight movement, turns, acceleration, and heading correction.", size=10.4)
    add_table(doc, ["Control idea", "How the robot uses it"], [
        ("Gyro", "Keeps heading stable during turns and straight driving."),
        ("Encoders", "Measure wheel rotation so distance moves can be repeated."),
        ("PID", "Adjusts motor power based on heading error."),
        ("Acceleration", "Smooth starts and stops reduce slip."),
    ], widths=[1.45, 4.9])
    add_p(doc, "Website result to highlight: Pybricks + PID correction achieved 94.2% consistency during testing.", size=12, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # Page 17
    add_page_title(doc, "Diagnostics", "Robot health before a run")
    add_editable_diagnostics(doc)
    doc.add_page_break()

    # Page 18
    add_page_title(doc, "Testing loop", "Rough draft, tune, refine")
    add_table(doc, ["Step", "What we do", "Evidence"], [
        ("1. Rough draft", "Prove the mission path can work.", "Early attachment photos and first run code."),
        ("2. Tune + debug", "Fix alignment, drift, timing, and attachment contact.", "PID values, heading prints, repeated test runs."),
        ("3. Refine", "Make the run consistent enough for competition.", "Final photos, color-coded tools, diagnostics checks."),
    ], widths=[1.25, 3.15, 2.0])
    add_bullets(doc, [
        "Captain: Isabella.",
        "Lead coder: Anthony.",
        "Lead builder: Kyle.",
        "Red side team: Brian, Owen, Sam.",
        "Blue side team: Cailey, Jonathan, Julisa, Jasper.",
    ])
    doc.add_page_break()

    # Page 19
    add_page_title(doc, "Judge reference", "Questions this book helps us answer")
    add_bullets(doc, [
        "How did your robot improve from earlier versions?",
        "Why did you choose these mission groups?",
        "Which attachments changed the most, and what did each change fix?",
        "How does your program keep the robot consistent?",
        "How do diagnostics help you know whether a failure is hardware or software?",
    ])
    add_table(doc, ["Spec", "Current answer"], [
        ("Platform", "LEGO SPIKE Prime with Pybricks"),
        ("Drive", "DriveBase using motors on ports A and E"),
        ("Attachment motors", "LAM on port B, RAM on port C"),
        ("Wheel diameter", "56 mm"),
        ("Default axle track", "122 mm, with per-run overrides"),
        ("Readiness checks", "Gyro drift, motor sync, turn accuracy, battery voltage"),
    ], widths=[1.65, 4.7])
    doc.add_page_break()

    # Page 20
    add_page_title(doc, "Closing", "Our design story in one sentence")
    add_p(doc, "We built a robot system where the physical base, attachments, strategy, code, and diagnostics all support the same goal: repeatable performance under competition pressure.", size=15, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    add_picture(doc, ROOT / "Robot Iterations/base robot iteration 6.2.jpg", width=5.8)
    add_caption(doc, "Final robot design reference photo.")

    # Appendix
    doc.add_page_break()
    add_page_title(doc, "Appendix", "Calibration notes from robot_base.py")
    add_code_box(doc, "Calibration guide", code_excerpt(robot_base, "STEP 1", 0, 18), "These notes document how wheel diameter, axle track, and heading PID are tuned.")

    doc.save(DOCX)
    return DOCX


def page_canvas():
    return Image.new("RGB", (1700, 2200), "white")


def draw_page_header(d, tag, title):
    d.rectangle((0, 0, 1700, 92), fill=hex_to_rgb(NAVY))
    d.text((90, 30), "AI FUTURE  |  ROBOT DESIGN TECHNICAL BOOK", fill="white", font=font(28, True))
    d.text((90, 130), tag.upper(), fill=hex_to_rgb(MUTED), font=font(25, True))
    d.text((90, 172), title, fill=hex_to_rgb(NAVY), font=font(56, True))
    d.line((90, 250, 1610, 250), fill=hex_to_rgb(BORDER), width=3)


def draw_footer(d, page_num):
    d.line((90, 2112, 1610, 2112), fill=hex_to_rgb(BORDER), width=2)
    d.text((90, 2135), "AI Future - Unearthed Robot Design", fill=hex_to_rgb(MUTED), font=font(22))
    d.text((1530, 2135), f"{page_num:02d}", fill=hex_to_rgb(MUTED), font=font(22, True))


def draw_body_text(d, x, y, text, size=30, color=INK, bold=False, width=1450, gap=10):
    return draw_wrapped(d, (x, y), text, font(size, bold), hex_to_rgb(color), width, gap)


def draw_bullet_list(d, x, y, items, width=1450, color=INK):
    for item in items:
        d.ellipse((x, y + 12, x + 13, y + 25), fill=hex_to_rgb(GREEN))
        y = draw_body_text(d, x + 34, y, item, size=27, color=color, width=width - 40, gap=8) + 8
    return y


def draw_table_image(d, x, y, headers, rows, col_widths, row_h=92):
    total = sum(col_widths)
    d.rounded_rectangle((x, y, x + total, y + row_h * (len(rows) + 1)), radius=18, fill=(255, 255, 255), outline=hex_to_rgb(BORDER), width=3)
    cx = x
    for i, h in enumerate(headers):
        d.rectangle((cx, y, cx + col_widths[i], y + row_h), fill=hex_to_rgb("#E8EEF5"))
        d.text((cx + 18, y + 28), h, fill=hex_to_rgb(NAVY), font=font(23, True))
        cx += col_widths[i]
    for r_i, row in enumerate(rows):
        cy = y + row_h * (r_i + 1)
        cx = x
        for c_i, val in enumerate(row):
            d.rectangle((cx, cy, cx + col_widths[c_i], cy + row_h), outline=hex_to_rgb(BORDER), width=1)
            draw_body_text(d, cx + 18, cy + 20, str(val), size=21, width=col_widths[c_i] - 36, gap=4)
            cx += col_widths[c_i]
    return y + row_h * (len(rows) + 1)


def place_image(canvas, path, box):
    x, y, w, h = box
    img = contain_fit(path, (w, h), fill=(255, 255, 255))
    canvas.paste(img, (x, y))


def code_card(d, x, y, title, code, note):
    d.rounded_rectangle((x, y, x + 1480, y + 430), radius=22, fill=hex_to_rgb("#F8FAFC"), outline=hex_to_rgb("#CBD5E1"), width=3)
    d.text((x + 26, y + 22), title, fill=hex_to_rgb("#2563EB"), font=font(27, True))
    yy = y + 72
    for line in code.splitlines()[:9]:
        d.text((x + 30, yy), line[:90], fill=hex_to_rgb("#0F172A"), font=font(20))
        yy += 28
    draw_body_text(d, x + 30, y + 350, note, size=21, color=MUTED, width=1400)
    return y + 460


def make_pdf_pages(assets):
    PROOF.mkdir(exist_ok=True)
    pages = []

    def save_page(img, n):
        draw_footer(ImageDraw.Draw(img), n)
        path = PROOF / f"page-{n:02d}.png"
        img.save(path, quality=94)
        pages.append(path)

    # Cover
    cover = cover_fit(assets["cover"], (1700, 2200))
    save_page(cover, 1)

    specs = [
        ("Judge quick summary", "What we engineered this season", "Our design goal was to make a robot the whole team could align, repair, explain, and trust under competition pressure.", None, [
            "Stable base: lower, cleaner, and easier to inspect than earlier concepts.",
            "Mission strategy: grouped nearby missions into short runs to reduce risk.",
            "Attachment system: active + active + passive mechanisms for strong multi-mission tools.",
            "Programming: Pybricks with gyro-enabled DriveBase, encoder distance, and tuned PID.",
            "Competition readiness: diagnostics for gyro drift, motor sync, turn accuracy, and battery voltage.",
        ]),
        ("Robot evolution", "From concept to reliable competition base", "The robot moved from broad concept work toward a compact base, cleaner motor layout, and attachment access that teammates could understand quickly.", ROOT / "initial concept this year.png", [
            "Earlier ideas helped us decide what was too large, fragile, or hard to reset.",
            "The final base favors repeatable alignment and easy attachment access.",
        ]),
        ("Current robot", "Anatomy of the final base", "The robot anatomy overlay connects the physical build to the code architecture and attachment system.", assets["robot"], [
            "Drive motors are initialized once in the shared base file.",
            "Attachment motors are separated into left and right mechanisms.",
            "The base is designed around repeatable launch alignment and fast swaps.",
        ]),
        ("Mission strategy", "Grouped runs reduce risk", "Colored paths show how each mission group moves from launch to targets.", assets["mission"], []),
        ("Mission models", "What each run is built to solve", "Mission model photos are grouped by run so judges can connect strategy to the table.", assets["mission_grid"], []),
    ]

    n = 2
    for tag, title, lead, img_path, bullets in specs:
        img = page_canvas()
        d = ImageDraw.Draw(img)
        draw_page_header(d, tag, title)
        y = draw_body_text(d, 90, 285, lead, size=30, color=INK, bold=False, width=1480) + 25
        if img_path:
            if tag == "Mission strategy":
                place_image(img, img_path, (110, 350, 1480, 1500))
                y = 1880
            elif tag == "Mission models":
                place_image(img, img_path, (90, y, 1520, 1320))
                y += 1360
            else:
                place_image(img, img_path, (90, y, 1520, 1080))
                y += 1120
        if bullets:
            y = draw_bullet_list(d, 110, y + 15, bullets, width=1450)
        save_page(img, n)
        n += 1

    # Run plan page
    img = page_canvas()
    d = ImageDraw.Draw(img)
    draw_page_header(d, "Run plan", "Current competition map")
    draw_body_text(d, 90, 285, "The plan favors shorter trips, repeatable alignment, and attachments that solve multiple models before returning home.", size=30)
    rows = [(r, m, launch, note) for r, m, _, launch, note in RUNS]
    draw_table_image(d, 90, 420, ["Run", "Missions", "Launch", "Design focus"], rows, [170, 270, 330, 760], row_h=112)
    save_page(img, n)
    n += 1

    for run in ATTACHMENTS:
        color = next(c for r, _, c, _, _ in RUNS if r == run)
        for page_idx, timeline_path in enumerate(assets[f"{run}_timeline"], start=1):
            img = page_canvas()
            d = ImageDraw.Draw(img)
            part = "versions 1-2" if page_idx == 1 else "versions 3-4"
            draw_page_header(d, "Attachment iteration", f"{run}: {part}")
            d.rectangle((90, 265, 1610, 335), fill=hex_to_rgb(color))
            d.text((120, 286), "Vertical run iteration page: fewer photos per sheet, much larger evidence images.", fill="white", font=font(27, True))
            place_image(img, timeline_path, (90, 350, 1520, 1660))
            save_page(img, n)
            n += 1

    # Attachment system
    img = page_canvas()
    d = ImageDraw.Draw(img)
    draw_page_header(d, "Attachment system", "Two active actions plus one passive action")
    draw_body_text(d, 90, 285, "Our strongest attachments combined powered control with simple passive motion. The active mechanisms gave precision, while passive geometry made setup and release faster.", size=31)
    draw_table_image(d, 90, 470, ["Design rule", "Why it helped"], [
        ("Active attachment", "Precise movement for lifts, drops, and controlled pushes."),
        ("Active attachment", "Second powered action lets one run solve more than one model."),
        ("Passive attachment", "No motor timing needed, so the action is faster and simpler."),
        ("Color coding", "Teammates can grab the right run tool instantly."),
    ], [420, 1110], row_h=130)
    draw_bullet_list(d, 110, 1120, ["Run A: all black pieces.", "Run B: blue, grey, and yellow pieces.", "Run C: black and yellow pieces.", "Run D: all black pieces.", "Run E: all blue pieces."])
    save_page(img, n)
    n += 1

    # Programming architecture
    img = page_canvas()
    d = ImageDraw.Draw(img)
    draw_page_header(d, "Programming", "Pybricks gives us direct control")
    place_image(img, assets["code"], (90, 285, 1520, 900))
    robot_base = read_code("Team Robot Code/robot_base.py")
    code_card(d, 110, 1245, "robot_base.py excerpt", code_excerpt(robot_base, "def make_robot", 0, 10), "The shared factory enables gyro, resets heading safely, and returns a DriveBase ready for each run.")
    save_page(img, n)
    n += 1

    # Code decisions
    img = page_canvas()
    d = ImageDraw.Draw(img)
    draw_page_header(d, "Code decisions", "Real run files, real calibration choices")
    run_a = read_code("Team Robot Code/Run_A_updated.py")
    run_d = read_code("Team Robot Code/Run_D.py")
    y = code_card(d, 110, 300, "Run_A_updated.py", code_excerpt(run_a, "robot.heading_control.pid", 3, 10), "Run A uses axle_track=101.6 and tuned heading PID because it solves a large mission chain.")
    code_card(d, 110, y + 35, "Run_D.py", code_excerpt(run_d, "robot.settings", 1, 8), "Run D uses axle_track=152 and slower settings for precision lift/drop work.")
    save_page(img, n)
    n += 1

    # Menu and match flow
    img = page_canvas()
    d = ImageDraw.Draw(img)
    draw_page_header(d, "Menu and match flow", "One simple selector for competition")
    menu = read_code("Team Robot Code/Menu.py")
    y = code_card(d, 110, 300, "Menu.py", code_excerpt(menu, "selected = hub_menu", 0, 18), "The center button is the stop button, the hub light gives status, and D opens diagnostics.")
    draw_bullet_list(d, 110, y + 30, ["The menu prevents hunting through files at the table.", "Each selected number imports exactly one run file.", "Diagnostics are available from the same interface as mission runs."])
    save_page(img, n)
    n += 1

    # Sensors
    img = page_canvas()
    d = ImageDraw.Draw(img)
    draw_page_header(d, "Sensors and PID", "Gyro heading plus encoder distance")
    draw_body_text(d, 90, 285, "The robot uses the gyro for heading and motor encoders for distance. Pybricks DriveBase lets us tune straight movement, turns, acceleration, and heading correction.", size=31)
    draw_table_image(d, 90, 470, ["Control idea", "How the robot uses it"], [
        ("Gyro", "Keeps heading stable during turns and straight driving."),
        ("Encoders", "Measure wheel rotation so distance moves can be repeated."),
        ("PID", "Adjusts motor power based on heading error."),
        ("Acceleration", "Smooth starts and stops reduce slip."),
    ], [350, 1180], row_h=135)
    d.rounded_rectangle((200, 1170, 1500, 1370), radius=28, fill=hex_to_rgb("#E8F5EF"), outline=hex_to_rgb(GREEN), width=4)
    d.text((290, 1228), "94.2% consistency", fill=hex_to_rgb(NAVY), font=font(68, True))
    d.text((300, 1310), "Pybricks + PID correction from website testing narrative", fill=hex_to_rgb(MUTED), font=font(27))
    save_page(img, n)
    n += 1

    # Diagnostics
    img = page_canvas()
    d = ImageDraw.Draw(img)
    draw_page_header(d, "Diagnostics", "Robot health before a run")
    place_image(img, assets["diagnostics"], (90, 300, 1520, 940))
    draw_bullet_list(d, 110, 1300, ["Diagnostics separate code tuning from hardware problems.", "Battery, cables, weak motors, and gyro drift can be caught before match pressure starts."])
    save_page(img, n)
    n += 1

    # Testing loop
    img = page_canvas()
    d = ImageDraw.Draw(img)
    draw_page_header(d, "Testing loop", "Rough draft, tune, refine")
    draw_table_image(d, 90, 315, ["Step", "What we do", "Evidence"], [
        ("1. Rough draft", "Prove the mission path can work.", "Early attachment photos and first run code."),
        ("2. Tune + debug", "Fix alignment, drift, timing, and attachment contact.", "PID values, heading prints, repeated runs."),
        ("3. Refine", "Make the run competition consistent.", "Final photos, color-coded tools, diagnostics."),
    ], [330, 650, 550], row_h=150)
    draw_bullet_list(d, 110, 920, ["Captain: Isabella.", "Lead coder: Anthony.", "Lead builder: Kyle.", "Red side team: Brian, Owen, Sam.", "Blue side team: Cailey, Jonathan, Julisa, Jasper."])
    save_page(img, n)
    n += 1

    # Reference
    img = page_canvas()
    d = ImageDraw.Draw(img)
    draw_page_header(d, "Judge reference", "Questions this book helps us answer")
    y = draw_bullet_list(d, 110, 310, [
        "How did your robot improve from earlier versions?",
        "Why did you choose these mission groups?",
        "Which attachments changed the most, and what did each change fix?",
        "How does your program keep the robot consistent?",
        "How do diagnostics help you know whether a failure is hardware or software?",
    ])
    draw_table_image(d, 90, y + 45, ["Spec", "Current answer"], [
        ("Platform", "LEGO SPIKE Prime with Pybricks"),
        ("Drive", "DriveBase using motors on ports A and E"),
        ("Attachment motors", "LAM on port B, RAM on port C"),
        ("Wheel diameter", "56 mm"),
        ("Default axle track", "122 mm, with per-run overrides"),
        ("Readiness checks", "Gyro drift, motor sync, turn accuracy, battery voltage"),
    ], [390, 1140], row_h=110)
    save_page(img, n)
    n += 1

    # Closing
    img = page_canvas()
    d = ImageDraw.Draw(img)
    draw_page_header(d, "Closing", "Our design story in one sentence")
    draw_body_text(d, 150, 350, "We built a robot system where the physical base, attachments, strategy, code, and diagnostics all support the same goal: repeatable performance under competition pressure.", size=45, color=NAVY, bold=True, width=1400, gap=14)
    place_image(img, ROOT / "Robot Iterations/base robot iteration 6.2.jpg", (260, 760, 1180, 850))
    save_page(img, n)
    n += 1

    # Appendix
    img = page_canvas()
    d = ImageDraw.Draw(img)
    draw_page_header(d, "Appendix", "Calibration notes from robot_base.py")
    code_card(d, 110, 310, "Calibration guide", code_excerpt(robot_base, "STEP 1", 0, 18), "These notes document how wheel diameter, axle track, and heading PID are tuned.")
    save_page(img, n)

    # Save multipage PDF from page proofs.
    pil_pages = [Image.open(p).convert("RGB") for p in pages]
    pil_pages[0].save(PDF, save_all=True, append_images=pil_pages[1:], resolution=200.0)

    # Contact sheets for visual QA.
    thumbs = [p.copy() for p in pil_pages]
    for thumb in thumbs:
        thumb.thumbnail((255, 330), Image.Resampling.LANCZOS)
    for sheet_idx in range(math.ceil(len(thumbs) / 12)):
        sheet = Image.new("RGB", (4 * 285, 3 * 375), "white")
        sd = ImageDraw.Draw(sheet)
        for i, thumb in enumerate(thumbs[sheet_idx * 12:(sheet_idx + 1) * 12]):
            x = (i % 4) * 285 + 15
            y = (i // 4) * 375 + 15
            sheet.paste(thumb, (x, y))
            sd.text((x, y + thumb.height + 8), f"Page {sheet_idx * 12 + i + 1}", fill=hex_to_rgb(MUTED), font=font(18, True))
        sheet.save(PROOF / f"contact_sheet_{sheet_idx + 1}.png", quality=94)
    return PDF


def main():
    assets = create_assets()
    build_docx(assets)
    make_pdf_pages(assets)
    print(DOCX)
    print(PDF)


if __name__ == "__main__":
    main()
