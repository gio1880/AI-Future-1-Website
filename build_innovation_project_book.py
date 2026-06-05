from pathlib import Path
import re
import textwrap

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Innovation Project Technical Book"
PROOF = OUT / "pdf_page_proofs"
ASSETS = OUT / "generated_assets"
DOCX = OUT / "AI Future - Innovation Project Technical Book.docx"
DOCX_MAPPING_UPDATE = OUT / "AI Future - Innovation Project Technical Book - Mapping Update.docx"
PDF = OUT / "AI Future - Innovation Project Technical Book.pdf"
PDF_MAPPING_UPDATE = OUT / "AI Future - Innovation Project Technical Book - Mapping Update.pdf"

NAVY = "#0B2545"
INK = "#111827"
MUTED = "#64748B"
LIGHT = "#F4F7FA"
BORDER = "#D9E2EC"
TEAL = "#00A6A6"
GREEN = "#2F9E6D"
GOLD = "#C28A13"
RED = "#C44545"


def hex_to_rgb(hex_color):
    h = hex_color.lstrip("#")
    return tuple(int(h[i:i + 2], 16) for i in (0, 2, 4))


def font(size=28, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def contain_fit(path, size, fill=(255, 255, 255)):
    try:
        img = ImageOps.exif_transpose(Image.open(path)).convert("RGB")
    except Exception:
        img = Image.new("RGB", size, fill)
        d = ImageDraw.Draw(img)
        d.text((20, 20), "Image unavailable", fill=hex_to_rgb(MUTED), font=font(22, True))
        return img
    img.thumbnail(size, Image.Resampling.LANCZOS)
    canvas = Image.new("RGB", size, fill)
    canvas.paste(img, ((size[0] - img.width) // 2, (size[1] - img.height) // 2))
    return canvas


def cover_fit(path, size):
    img = ImageOps.exif_transpose(Image.open(path)).convert("RGB")
    return ImageOps.fit(img, size, method=Image.Resampling.LANCZOS, centering=(0.5, 0.5))


def draw_wrapped(draw, xy, text, fnt, fill, width, line_gap=6):
    x, y = xy
    avg = max(1, int(width / (fnt.size * 0.55)))
    lines = []
    for para in str(text).splitlines():
        lines.extend(textwrap.wrap(para, avg) or [""])
    for line in lines:
        draw.text((x, y), line, fill=fill, font=fnt)
        y += fnt.size + line_gap
    return y


def round_rect(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def style_run(run, size=11, bold=False, color=INK, font_name="Aptos"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(*hex_to_rgb(color))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.lstrip("#"))
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color.lstrip("#"))


def add_p(doc, text="", size=10.5, color=INK, bold=False, align=None, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        style_run(r, size=size, bold=bold, color=color)
    return p


def add_h(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    style_run(r, size=18 if level == 1 else 13, bold=True, color=NAVY if level == 1 else "#1F4D78")
    return p


def add_label(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    style_run(r, size=8.5, bold=True, color=color)
    return p


def add_page_title(doc, tag, title, summary=None):
    add_label(doc, tag)
    add_h(doc, title, 1)
    if summary:
        add_p(doc, summary, size=10.4, color=MUTED, after=5)


def add_picture(doc, path, width=6.9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(width))
    return p


def add_caption(doc, text):
    return add_p(doc, text, size=8.2, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        r = p.add_run(item)
        style_run(r, size=10.1, color=INK)


def add_cell_text(cell, text, size=8.8, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        if widths:
            cell.width = Inches(widths[i])
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(cell)
        add_cell_text(cell, header, size=8.4, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if widths:
                cells[i].width = Inches(widths[i])
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            add_cell_text(cells[i], str(value), size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def set_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("AI Future - Innovation Project Technical Book")
    style_run(r, size=8.2, color=MUTED)


def read_code(path):
    return (ROOT / path).read_text(encoding="utf-8", errors="replace").replace("â‰ˆ", "~")


def code_excerpt(text, contains, before=1, after=12):
    lines = text.splitlines()
    idx = next((i for i, line in enumerate(lines) if contains in line), 0)
    return "\n".join(lines[max(0, idx - before):min(len(lines), idx + after)])


def add_code_box(doc, title, code, note):
    add_label(doc, title, "#2563EB")
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.9)
    set_cell_shading(cell, "F8FAFC")
    set_cell_border(cell, "CBD5E1")
    p = cell.paragraphs[0]
    r = p.add_run(code)
    style_run(r, size=7.7, color="#0F172A", font_name="Consolas")
    add_p(doc, note, size=8.5, color=MUTED)


def add_mapping_technical_pages(doc):
    add_page_title(doc, "Current mapping system", "How our tunnel scan works now", "Our current mapping prototype uses the camera to guide overlapping photos, then stitches those photos into a wider room/tunnel view.")
    add_table(doc, ["Step", "Code/file", "What happens"], [
        ("1. Open camera", "guided_room_scan.py", "Uses OpenCV VideoCapture at 640 x 480 and 15 FPS."),
        ("2. Detect image motion", "frame_motion()", "ORB finds visual features in the previous capture and current frame."),
        ("3. Match features", "BFMatcher", "Feature matches estimate how far the camera has turned between shots."),
        ("4. Decide capture timing", "guide_text()", "The guide says turn right, move slower, or READY based on shift ratio and match count."),
        ("5. Stitch output", "stitch_room_snapshots.py", "OpenCV Stitcher combines overlapping snapshots into a panorama."),
    ], widths=[1.25, 1.75, 3.4])
    add_code_box(
        doc,
        "guided_room_scan.py: motion logic",
        "MIN_SHIFT_RATIO = 0.14\nMAX_SHIFT_RATIO = 0.42\nMIN_GOOD_MATCHES = 25\n\nmedian_shift = median(previous_x - current_x)\nshift_ratio = median_shift / frame_width",
        "This is the current camera-turn logic. A good scan needs enough shared visual features and the right amount of overlap between photos.",
    )
    doc.add_page_break()

    add_page_title(doc, "Mapping math", "From sensor readings to map points", "Before the camera-first system, we tested distance readings and converted them into point-cloud coordinates. This helped us understand the math behind mapping a tunnel.")
    add_table(doc, ["Math idea", "Formula", "Meaning"], [
        ("Angle step", "angle_deg = angle_index x 15", "Each scan step estimates the rover/camera direction."),
        ("Front sensor point", "x = r cos(theta), y = r sin(theta)", "A distance reading becomes a 2D point in front of the rover."),
        ("Back sensor point", "theta_back = theta + 180", "Back readings are plotted behind the rover."),
        ("No-hit filter", "ignore readings >= 2000 mm", "Large values are treated as open space or no detected wall."),
        ("Robot origin", "(0, 0)", "The rover starts as the center point of the local map."),
    ], widths=[1.45, 2.15, 2.8])
    code = read_code("innovationextension.py")
    add_code_box(
        doc,
        "innovationextension.py: coordinate conversion",
        code_excerpt(code, "front_x.append", 8, 18),
        "This baseline math is useful for explaining how a tunnel wall can be represented as plotted points.",
    )
    doc.add_page_break()

    add_page_title(doc, "Current output", "What our map looks like today", "Right now our strongest output is visual: a contact sheet of captured frames and a stitched panorama of the actual STEM lab.")
    add_photo_table(doc, "Editable mapping output evidence", [
        ("room_model/room_contact_sheet.jpg", "Contact sheet: the individual overlapping frames captured during the guided scan."),
        ("room_model/room_panorama.jpg", "Panorama: the stitched result that shows the actual STEM lab as one wider map-like view."),
        ("Rover Iterations/distance sensor mapping results.png", "Distance baseline: early point-cloud result from sensor readings."),
    ])
    add_bullets(doc, [
        "Current strength: shows real visual context instead of only dots or distances.",
        "Current limitation: panorama output is not yet a measured floor plan.",
        "Why it matters: archaeologists need context, surfaces, obstacles, and room shape, not just one distance number.",
    ])
    doc.add_page_break()

    add_page_title(doc, "Future tunnel mapper", "What the next mapping system would look like", "Our future version would combine camera images, pose estimates, and a tunnel map that updates as the rover moves.")
    add_table(doc, ["Future feature", "Planned logic", "Why it improves the project"], [
        ("Keyframes", "Save important camera views when overlap is good.", "Reduces blurry or duplicate frames."),
        ("Pose estimate", "Track turn angle and distance traveled between frames.", "Places each image in a tunnel path instead of only stitching side-by-side."),
        ("Obstacle map", "Mark blocked/open areas from camera and sensor cues.", "Helps archaeologists decide where the rover can safely continue."),
        ("Loop check", "Recognize when the rover sees a place again.", "Reduces map drift during longer tunnel scans."),
        ("VR export", "Convert panorama/map sections into a walkthrough.", "Lets judges, museums, or students explore the site safely."),
    ], widths=[1.45, 2.5, 2.5])
    add_code_box(
        doc,
        "Future pseudocode",
        "for each camera frame:\n    features = detect_features(frame)\n    pose = estimate_motion(previous_frame, frame)\n    if overlap_is_good(pose, features):\n        save_keyframe(frame, pose)\n        update_tunnel_map(keyframes)\n        mark_open_space_and_obstacles()\nexport_panorama_and_vr_scene()",
        "This is the next version of the logic: not just taking pictures, but building a tunnel map that understands movement and obstacles.",
    )
    doc.add_page_break()


def add_photo_table(doc, title, photos):
    add_label(doc, title, TEAL)
    table = doc.add_table(rows=len(photos), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (rel, caption) in enumerate(photos):
        img_cell, text_cell = table.rows[idx].cells
        for cell in (img_cell, text_cell):
            set_cell_border(cell, "CBD5E1")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        img_cell.width = Inches(3.35)
        text_cell.width = Inches(3.35)
        img_cell.text = ""
        p = img_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(ROOT / rel), width=Inches(3.08))
        add_cell_text(text_cell, caption, size=9.2, color=INK, bold=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_docx():
    OUT.mkdir(exist_ok=True)
    ASSETS.mkdir(exist_ok=True)
    doc = Document()
    set_page(doc)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(10.5)

    add_picture(doc, ROOT / "Rover Iterations/recent innvoation pictures/10-innovation-rover-front-view-sensor-mast.JPG", width=7.15)
    add_p(doc, "AI Future - Innovation Project", size=22, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Safer Mapping for Fragile Archaeological Sites", size=15, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    add_page_title(doc, "Judge quick summary", "What problem did we solve?", "Archaeologists need safer ways to document dark, unstable, tight spaces without putting people or fragile cultural heritage at risk.")
    add_bullets(doc, [
        "Solution: a rover concept for remote observation and mapping.",
        "Current upgrade: wooden rover platform for stronger protection and durability.",
        "Programming upgrade: camera-first mapping logic instead of relying on distance sensors as the primary method.",
        "Impact: safer exploration, richer documentation, and future VR walkthroughs for education and museums.",
    ])
    add_table(doc, ["Area", "Evidence judges can inspect"], [
        ("Research", "Museum visit, documentaries, archaeology problem framing, expert feedback"),
        ("Prototypes", "Cardboard, LEGO, CAD, 3D printed/tread ideas, wooden rover"),
        ("Testing", "Distance sensor baseline, camera mapping screenshots, stitched room panorama"),
        ("Feedback", "Grace Howe, Rawlins, teachers, mentors"),
    ], widths=[1.4, 5.0])
    doc.add_page_break()

    add_page_title(doc, "Identify", "We started with real archaeology risks")
    add_picture(doc, ROOT / "Rover Iterations/museum trip.png", width=6.8)
    add_caption(doc, "Museum research helped the team understand artifact fragility and preservation needs.")
    add_table(doc, ["Discovery method", "What we learned"], [
        ("Observe", "Artifacts are fragile and documentation must avoid accidental damage."),
        ("Research", "Dark, unstable, hard-to-reach spaces create safety and mapping challenges."),
        ("Interview", "Experts pushed us toward camera height, blind-spot reduction, and modular repairability."),
        ("Define", "A rover could reduce risk by sending sensors/cameras where people should not go first."),
    ], widths=[1.5, 4.9])
    doc.add_page_break()

    add_page_title(doc, "Solution", "A rover for safer mapping")
    add_picture(doc, ROOT / "Rover Iterations/recent innvoation pictures/10-innovation-rover-front-view-sensor-mast.JPG", width=6.9)
    add_caption(doc, "Recent rover version with front sensor mast and updated wood-panel protection.")
    add_bullets(doc, [
        "Remote observation helps reduce risk before people enter a fragile site.",
        "Camera placement is designed to capture walls, ceilings, obstacles, and context.",
        "The rover is modular so parts can be repaired or replaced after field-like tests.",
    ])
    doc.add_page_break()

    add_page_title(doc, "Create", "Prototype evolution")
    add_photo_table(doc, "Editable prototype timeline", [
        ("Rover Iterations/cardboard models.png", "Cardboard models: fast layout choices for size, wiring paths, and camera location."),
        ("Rover Iterations/innovation lego prototype.jpg", "LEGO prototype: movement and stability testing with a physical model."),
        ("Rover Iterations/recent innvoation pictures/02-innovation-rover-wood-panel-side-view.JPG", "Recent wood-panel rover: stronger protective side panels and updated sensor placement."),
        ("Rover Iterations/recent innvoation pictures/06-team-holding-innovation-rover-prototype.JPG", "Recent team-held prototype: current physical version ready for judge explanation."),
    ])
    doc.add_page_break()

    add_page_title(doc, "Iterate", "Feedback changed the design")
    add_table(doc, ["Feedback source", "What changed"], [
        ("Grace Howe, archaeologist", "Raised/repositioned camera to reduce blind spots and capture useful wall/ceiling detail."),
        ("Rawlins, event organizer", "Pushed modularity so repairs and field adjustments would be easier."),
        ("Teachers and mentors", "Noted wheel limitations on uneven terrain; team researched treads and protective plates."),
        ("Team testing", "Moved from distance-sensor-first mapping to camera-based observation and mapping."),
    ], widths=[1.9, 4.5])
    add_picture(doc, ROOT / "Rover Iterations/gracehowepicture.jpg", width=3.9)
    add_caption(doc, "Expert feedback was treated as part of the engineering loop, not just a presentation detail.")
    doc.add_page_break()

    add_page_title(doc, "Latest upgrade", "Wooden rover + camera mapping")
    add_photo_table(doc, "Editable latest-build evidence", [
        ("Rover Iterations/recent innvoation pictures/10-innovation-rover-front-view-sensor-mast.JPG", "Recent front view: sensor mast and rover layout for mapping demonstrations."),
        ("Rover Iterations/recent innvoation pictures/11-innovation-rover-table-demo-with-globe.JPG", "Recent table demo: rover shown with globe/terrain context for explaining mapping use."),
        ("Rover Iterations/recent innvoation pictures/15-electronics-rover-control-board-closeup.JPG", "Electronics prototype: control board and wiring evidence for sensor/mapping experiments."),
        ("Rover Iterations/recent innvoation pictures/23-electronics-rover-front-view.JPG", "Electronics rover front view: ultrasonic sensor setup and updated drive platform."),
        ("room_model/room_panorama.jpg", "Python/OpenCV stitched panorama from guided camera scan."),
        ("room_model/room_contact_sheet.jpg", "Contact sheet showing captured frames used for the room model."),
    ])
    doc.add_page_break()

    add_page_title(doc, "Actual STEM lab scan", "Our real room panorama", "This panorama shows our actual STEM lab scan from the guided camera mapping prototype.")
    add_picture(doc, ROOT / "room_model/room_panorama.jpg", width=7.15)
    add_caption(doc, "Actual STEM lab panorama stitched from overlapping camera frames.")
    add_bullets(doc, [
        "The scan was captured in our real STEM lab workspace.",
        "Overlapping frames were stitched together into one wider room model.",
        "This is evidence of the mapping workflow we want to improve for archaeology and museum use.",
    ])
    doc.add_page_break()

    add_page_title(doc, "Programming", "Camera-first observation + mapping logic")
    add_table(doc, ["Programming idea", "Current status"], [
        ("Observe surroundings", "Camera captures visual context instead of only distance readings."),
        ("Classify open space/obstacles", "Website demo and Python tools support guided scans."),
        ("Build map output", "Python OpenCV stitching combines overlapping frames into a panorama."),
        ("Debug and improve", "Saved photos/contact sheets help compare mapping quality between tests."),
    ], widths=[2.0, 4.4])
    add_code_box(doc, "guided_room_scan.py command", ".\\.venv\\Scripts\\python.exe guided_room_scan.py", "This command runs the guided Python scan referenced by the website.")
    doc.add_page_break()

    add_mapping_technical_pages(doc)

    add_page_title(doc, "Early baseline", "Distance sensor mapping taught us what to improve")
    add_picture(doc, ROOT / "Rover Iterations/distance sensor mapping results.png", width=6.8)
    add_caption(doc, "Early distance-sensor mapping results before shifting to camera-based mapping.")
    code = read_code("innovationextension.py")
    add_code_box(doc, "innovationextension.py excerpt", code_excerpt(code, "ANGLE_STEP_DEG", 2, 16), "The baseline script converted front/back distance readings into point-cloud coordinates.")
    doc.add_page_break()

    add_page_title(doc, "Impact", "Protect people and history")
    add_photo_table(doc, "Why the project matters", [
        ("Rover Iterations/recent innvoation pictures/08-innovation-rover-front-sensor-closeup.JPG", "Recent sensor closeup: camera/sensor placement is central to mapping tight spaces."),
        ("Rover Iterations/recent innvoation pictures/24-electronics-rover-front-angle-view.JPG", "Recent electronics rover: continued testing with sensing hardware and mobile platform."),
        ("Rover Iterations/vr experience.jpg", "Future direction: turn mapping output into VR so museums and students can explore safely."),
    ])
    doc.add_page_break()

    add_page_title(doc, "Judge reference", "Questions this book helps answer")
    add_bullets(doc, [
        "What real-world problem did you identify?",
        "How did expert feedback change your design?",
        "What did each prototype teach the team?",
        "Why did you move from distance sensors to camera mapping?",
        "How could this help archaeologists, museums, or students?",
    ])
    add_table(doc, ["Spec", "Current answer"], [
        ("Problem", "Dark, unstable, tight archaeological spaces are risky to map manually."),
        ("Solution", "Remote rover concept for camera-based mapping and documentation."),
        ("Latest build", "Wooden rover platform with camera mapping workflow."),
        ("Software", "Python/OpenCV guided room scan and early point-cloud code."),
        ("Future", "Better low-light capture and smoother VR walkthroughs."),
    ], widths=[1.35, 5.0])

    try:
        doc.save(DOCX)
        return DOCX
    except OSError:
        doc.save(DOCX_MAPPING_UPDATE)
        return DOCX_MAPPING_UPDATE


def page_canvas():
    return Image.new("RGB", (1700, 2200), "white")


def draw_page_header(d, tag, title):
    d.rectangle((0, 0, 1700, 92), fill=hex_to_rgb(NAVY))
    d.text((90, 30), "AI FUTURE  |  INNOVATION PROJECT TECHNICAL BOOK", fill="white", font=font(28, True))
    d.text((90, 130), tag.upper(), fill=hex_to_rgb(MUTED), font=font(25, True))
    d.text((90, 172), title, fill=hex_to_rgb(NAVY), font=font(54, True))
    d.line((90, 250, 1610, 250), fill=hex_to_rgb(BORDER), width=3)


def draw_footer(d, n):
    d.line((90, 2112, 1610, 2112), fill=hex_to_rgb(BORDER), width=2)
    d.text((90, 2135), "AI Future - Innovation Project", fill=hex_to_rgb(MUTED), font=font(22))
    d.text((1530, 2135), f"{n:02d}", fill=hex_to_rgb(MUTED), font=font(22, True))


def draw_text(d, x, y, text, size=30, color=INK, bold=False, width=1450):
    return draw_wrapped(d, (x, y), text, font(size, bold), hex_to_rgb(color), width, 8)


def draw_bullets(d, x, y, items):
    for item in items:
        d.ellipse((x, y + 13, x + 14, y + 27), fill=hex_to_rgb(TEAL))
        y = draw_text(d, x + 36, y, item, size=28, width=1400) + 10
    return y


def draw_table_box(d, x, y, title, rows, accent=TEAL):
    d.rounded_rectangle((x, y, x + 1500, y + 120 + len(rows) * 125), radius=22, fill=hex_to_rgb(LIGHT), outline=hex_to_rgb(BORDER), width=3)
    d.rectangle((x, y, x + 1500, y + 72), fill=hex_to_rgb(accent))
    d.text((x + 28, y + 20), title, fill="white", font=font(29, True))
    yy = y + 98
    for label, body in rows:
        d.text((x + 34, yy), label, fill=hex_to_rgb(NAVY), font=font(26, True))
        draw_text(d, x + 360, yy, body, size=23, width=1050)
        yy += 125
    return y + 120 + len(rows) * 125


def draw_code_panel(d, x, y, title, code, note):
    d.rounded_rectangle((x, y, x + 1500, y + 520), radius=22, fill=(248, 250, 252), outline=hex_to_rgb("#CBD5E1"), width=3)
    d.text((x + 28, y + 24), title, fill=hex_to_rgb("#2563EB"), font=font(28, True))
    yy = y + 78
    for line in code.splitlines()[:10]:
        d.text((x + 32, yy), line[:92], fill=hex_to_rgb("#0F172A"), font=font(21))
        yy += 31
    draw_text(d, x + 32, y + 420, note, size=22, color=MUTED, width=1400)
    return y + 550


def place_image(canvas, path, box):
    x, y, w, h = box
    canvas.paste(contain_fit(path, (w, h), fill=(255, 255, 255)), (x, y))


def make_pdf_pages():
    PROOF.mkdir(exist_ok=True)
    pages = []

    def save_page(img, n):
        draw_footer(ImageDraw.Draw(img), n)
        path = PROOF / f"page-{n:02d}.png"
        img.save(path, quality=94)
        pages.append(path)

    cover = cover_fit(ROOT / "Rover Iterations/recent innvoation pictures/10-innovation-rover-front-view-sensor-mast.JPG", (1700, 2200))
    d = ImageDraw.Draw(cover, "RGBA")
    d.rectangle((0, 0, 1700, 2200), fill=(255, 255, 255, 65))
    d.rectangle((0, 1320, 1700, 2200), fill=(255, 255, 255, 235))
    d.rectangle((0, 0, 1700, 120), fill=hex_to_rgb(NAVY) + (238,))
    d.text((90, 38), "AI FUTURE  |  UNEARTHED SEASON", fill="white", font=font(38, True))
    d.text((90, 1415), "Innovation Project", fill=hex_to_rgb(NAVY), font=font(92, True))
    d.text((94, 1530), "Safer Mapping for Fragile Archaeological Sites", fill=hex_to_rgb(NAVY), font=font(46, True))
    d.text((94, 1610), "Rover design, camera mapping, expert feedback, and future VR", fill=hex_to_rgb(MUTED), font=font(32))
    save_page(cover, 1)

    specs = [
        ("Judge quick summary", "What problem did we solve?", "Archaeologists need safer ways to document dark, unstable, tight spaces without putting people or fragile cultural heritage at risk.", ROOT / "Rover Iterations/recent innvoation pictures/06-team-holding-innovation-rover-prototype.JPG", ["Solution: remote rover concept for observation and mapping.", "Current upgrade: recent wood-panel rover platform with sensor mast.", "Programming upgrade: camera-first mapping instead of distance-sensor-first mapping.", "Impact: safer exploration and future VR walkthroughs."]),
        ("Identify", "Real archaeology risks", "We researched artifact fragility, difficult field conditions, dark caves, and expert concerns before choosing our problem.", ROOT / "Rover Iterations/museum trip.png", ["Museum visit helped us understand preservation.", "Expert feedback shaped camera placement and modularity."]),
        ("Solution", "A rover for safer mapping", "The rover sends cameras and mapping tools into spaces people should not enter first.", ROOT / "Rover Iterations/recent innvoation pictures/10-innovation-rover-front-view-sensor-mast.JPG", ["Stable movement.", "Camera and mapping focus.", "Repairable modular platform."]),
        ("Prototype evolution", "Physical + digital versions", "Each prototype answered a different engineering question.", ROOT / "Rover Iterations/recent innvoation pictures/02-innovation-rover-wood-panel-side-view.JPG", ["Cardboard: quick layout.", "LEGO: movement test.", "Wood-panel rover: stronger current platform.", "Electronics rover: sensor experiments."]),
        ("Latest upgrade", "Recent rover + camera mapping", "The current workflow combines the recent rover build with camera-based mapping output.", ROOT / "Rover Iterations/recent innvoation pictures/11-innovation-rover-table-demo-with-globe.JPG", ["Recent table demo version.", "OpenCV stitched room model.", "Camera logic gives richer context than distance-only readings."]),
        ("Actual STEM lab scan", "Our real room panorama", "This is the actual STEM lab panorama produced by the guided camera mapping prototype.", ROOT / "room_model/room_panorama.jpg", ["Captured from the team workspace.", "Stitched from overlapping camera frames.", "Shows the mapping workflow working on a real indoor space."]),
        ("Impact", "Protect people and history", "The project reduces risk while helping archaeologists and museums document and share fragile places.", ROOT / "Rover Iterations/vr experience.jpg", ["Safer pre-entry documentation.", "Better visual context.", "Future VR education experience."]),
    ]
    n = 2
    for tag, title, lead, img_path, bullets in specs:
        img = page_canvas()
        d = ImageDraw.Draw(img)
        draw_page_header(d, tag, title)
        y = draw_text(d, 90, 290, lead, size=31, width=1500) + 22
        place_image(img, img_path, (100, y, 1500, 1060))
        y += 1110
        draw_bullets(d, 115, y, bullets)
        save_page(img, n)
        n += 1

    # Editable-style evidence pages for PDF.
    evidence_sets = [
        ("Create", "Prototype evidence", [("Rover Iterations/recent innvoation pictures/02-innovation-rover-wood-panel-side-view.JPG", "Recent wood-panel rover"), ("Rover Iterations/recent innvoation pictures/06-team-holding-innovation-rover-prototype.JPG", "Team-held current prototype"), ("Rover Iterations/recent innvoation pictures/23-electronics-rover-front-view.JPG", "Electronics rover front view")]),
        ("Programming", "Mapping evidence", [("Rover Iterations/distance sensor mapping results.png", "Early distance baseline"), ("room_model/room_contact_sheet.jpg", "Captured camera frames"), ("room_model/room_panorama.jpg", "Stitched panorama output")]),
    ]
    for tag, title, photos in evidence_sets:
        img = page_canvas()
        d = ImageDraw.Draw(img)
        draw_page_header(d, tag, title)
        y = 310
        for rel, cap in photos:
            round_rect(d, (100, y, 1600, y + 520), 24, (248, 250, 252), hex_to_rgb(BORDER), 3)
            place_image(img, ROOT / rel, (130, y + 30, 760, 460))
            draw_text(d, 930, y + 70, cap, size=38, color=NAVY, bold=True, width=560)
            y += 560
        save_page(img, n)
        n += 1

    technical_pages = [
        (
            "Current mapping system",
            "How our tunnel scan works now",
            "The current prototype guides overlapping camera captures, then stitches the saved frames into a wider visual map.",
            [
                ("Open camera", "OpenCV reads 640 x 480 frames at 15 FPS."),
                ("Detect motion", "ORB feature points compare the previous capture to the current frame."),
                ("Guide capture", "Shift ratio and match count decide turn right, move slower, or READY."),
                ("Stitch output", "OpenCV Stitcher creates the panorama and a contact sheet."),
            ],
            "MIN_SHIFT_RATIO = 0.14\nMAX_SHIFT_RATIO = 0.42\nMIN_GOOD_MATCHES = 25\n\nmedian_shift = median(previous_x - current_x)\nshift_ratio = median_shift / frame_width",
            "Good tunnel scans need enough shared visual features and the right amount of overlap.",
        ),
        (
            "Mapping math",
            "From readings to map points",
            "The early distance-sensor version helped us explain the math of turning tunnel measurements into plotted points.",
            [
                ("Angle", "angle_deg = angle_index x 15"),
                ("Front point", "x = r cos(theta), y = r sin(theta)"),
                ("Back point", "theta_back = theta + 180"),
                ("No hit", "Ignore readings >= 2000 mm."),
            ],
            "a_deg = angle_index * ANGLE_STEP_DEG\nth = radians(a_deg)\nx = distance_mm * cos(th)\ny = distance_mm * sin(th)\n\n# back sensor uses the opposite direction\nth_back = radians(a_deg + 180)",
            "This converts rover-relative measurements into a local 2D point cloud.",
        ),
        (
            "Future tunnel mapper",
            "What the next system would look like",
            "The future mapper would combine images, pose estimates, obstacle marks, and VR export.",
            [
                ("Keyframes", "Save useful views only when overlap and image quality are good."),
                ("Pose estimate", "Track turn angle and distance traveled between frames."),
                ("Obstacle map", "Mark open and blocked areas as the rover advances."),
                ("VR export", "Turn the scan into an explorable tunnel walkthrough."),
            ],
            "for each camera frame:\n    features = detect_features(frame)\n    pose = estimate_motion(previous_frame, frame)\n    if overlap_is_good(pose, features):\n        save_keyframe(frame, pose)\n        update_tunnel_map(keyframes)\n        mark_open_space_and_obstacles()\nexport_panorama_and_vr_scene()",
            "This is the roadmap from camera demo to a real tunnel mapping system.",
        ),
    ]
    for tag, title, lead, rows, code, note in technical_pages:
        img = page_canvas()
        d = ImageDraw.Draw(img)
        draw_page_header(d, tag, title)
        y = draw_text(d, 90, 290, lead, size=31, width=1500) + 35
        y = draw_table_box(d, 90, y, "Logic summary", rows) + 35
        draw_code_panel(d, 90, y, "Code / math view", code, note)
        save_page(img, n)
        n += 1

    img = page_canvas()
    d = ImageDraw.Draw(img)
    draw_page_header(d, "Judge reference", "Questions this book helps answer")
    draw_bullets(d, 120, 330, [
        "What real-world problem did you identify?",
        "How did expert feedback change your design?",
        "What did each prototype teach the team?",
        "Why did you move from distance sensors to camera mapping?",
        "How could this help archaeologists, museums, or students?",
    ])
    save_page(img, n)

    pil_pages = [Image.open(p).convert("RGB") for p in pages]
    try:
        pil_pages[0].save(PDF, save_all=True, append_images=pil_pages[1:], resolution=200.0)
        return PDF
    except OSError:
        pil_pages[0].save(PDF_MAPPING_UPDATE, save_all=True, append_images=pil_pages[1:], resolution=200.0)
        return PDF_MAPPING_UPDATE


def main():
    OUT.mkdir(exist_ok=True)
    ASSETS.mkdir(exist_ok=True)
    build_docx()
    make_pdf_pages()
    print(DOCX)
    print(PDF)


if __name__ == "__main__":
    main()
